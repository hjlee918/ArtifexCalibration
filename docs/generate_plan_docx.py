"""
Generate the LG OLED Calibration App implementation plan as a .docx file.
Run: python3 docs/generate_plan_docx.py
Output: docs/LG_OLED_Calibration_App_Implementation_Plan.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).parent
OUTPUT_PATH = DOCS_DIR / "LG_OLED_Calibration_App_Implementation_Plan.docx"

PLAN_FILES = [
    ("Project Overview", DOCS_DIR / "superpowers/specs/2026-04-23-project-overview.md"),
    ("Sub-project 1: TV Control Layer", DOCS_DIR / "superpowers/plans/2026-04-16-subproject1-tv-control-layer.md"),
    ("Sub-project 2: LUT Upload Pipeline", DOCS_DIR / "superpowers/plans/2026-04-23-subproject2-lut-upload-pipeline.md"),
    ("Sub-project 3: Measurement Workflow", DOCS_DIR / "superpowers/plans/2026-04-23-subproject3-measurement-workflow.md"),
]

# Color scheme
COLOR_HEADING1 = RGBColor(0x1A, 0x73, 0xE8)   # Blue
COLOR_HEADING2 = RGBColor(0x0D, 0x47, 0xA1)   # Dark blue
COLOR_HEADING3 = RGBColor(0x15, 0x65, 0xC0)   # Medium blue
COLOR_CODE_BG  = RGBColor(0xF5, 0xF5, 0xF5)   # Light gray
COLOR_CODE_FG  = RGBColor(0x1A, 0x1A, 0x2E)   # Dark
COLOR_TASK_BG  = RGBColor(0xE3, 0xF2, 0xFD)   # Light blue
COLOR_ACCENT   = RGBColor(0x4F, 0xC3, 0xF7)   # Cyan accent


def set_cell_background(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def setup_styles(doc):
    """Configure document styles."""
    # Normal
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_HEADING1
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_HEADING2
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_HEADING3
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # Code style
    if 'Code' not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles['Code']
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.left_indent = Inches(0.2)


def add_title_page(doc):
    """Add a styled title page."""
    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(80)
    run = title_para.add_run("LG OLED Calibration App")
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING1

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle_para.add_run("Complete Implementation Plan")
    sub.font.name = 'Calibri'
    sub.font.size = Pt(18)
    sub.font.color.rgb = COLOR_HEADING2

    doc.add_paragraph()

    # Project description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d = desc.add_run(
        "A Calman-equivalent display calibration application for LG OLED TVs\n"
        "2021–2026 model years · macOS · Python + PyQt6"
    )
    d.font.name = 'Calibri'
    d.font.size = Pt(12)
    d.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    d.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Sub-project summary table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headers = ["Sub-project", "Scope", "Status"]
    rows_data = [
        ("1 — TV Control Layer",
         "TV discovery, SSAP connection, Expert menu read/write,\nKeychain pairing, model detection 2021–2026",
         "Plan written\nNot yet implemented"),
        ("2 — LUT Upload Pipeline",
         "1D/3D LUT parsing & upload, Dolby Vision\nconfig upload, LUT management UI",
         "Plan written\nNot yet implemented"),
        ("3 — Measurement Workflow",
         "ArgyllCMS meter integration, iTPG + PGenerator,\npatch sequences, CGATS storage, LUT generation",
         "Plan written\nNot yet implemented"),
    ]

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_background(cell, '1A73E8')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, (sp, scope, status) in enumerate(rows_data, start=1):
        row = table.rows[row_idx]
        cells = row.cells

        set_cell_background(cells[0], 'E3F2FD' if row_idx % 2 == 1 else 'FFFFFF')
        set_cell_background(cells[1], 'E3F2FD' if row_idx % 2 == 1 else 'FFFFFF')
        set_cell_background(cells[2], 'E3F2FD' if row_idx % 2 == 1 else 'FFFFFF')

        r0 = cells[0].paragraphs[0].add_run(sp)
        r0.font.name = 'Calibri'
        r0.font.bold = True
        r0.font.size = Pt(9.5)

        r1 = cells[1].paragraphs[0].add_run(scope)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9)

        r2 = cells[2].paragraphs[0].add_run(status)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9)
        r2.font.italic = True

    doc.add_paragraph()

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run("Generated: April 23, 2026")
    dr.font.name = 'Calibri'
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    add_page_break(doc)


def add_toc_entry(doc, number, title, page_hint=""):
    """Add a table of contents entry."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(f"{number}  {title}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    if number.count('.') == 0:
        run.font.bold = True
        run.font.color.rgb = COLOR_HEADING1


def add_toc(doc):
    """Add table of contents page."""
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = COLOR_HEADING1

    toc_entries = [
        ("1", "Project Overview"),
        ("1.1", "Goals and Hardware Targets"),
        ("1.2", "LG OLED Model / Chip / webOS Reference (2021–2026)"),
        ("1.3", "Full System Architecture"),
        ("1.4", "Calibration Workflow (End-to-End)"),
        ("1.5", "Technology Stack"),
        ("2", "Sub-project 1: TV Control Layer"),
        ("2.1", "Architecture and Components"),
        ("2.2", "Task 1: Project Scaffold"),
        ("2.3", "Task 2: TVSettingsSnapshot"),
        ("2.4", "Task 3: Discovery Service"),
        ("2.5", "Task 4: Keychain Utility"),
        ("2.6", "Task 5: Connection Manager"),
        ("2.7", "Task 6: LGTVSettings Extension"),
        ("2.8", "Task 7: PyQt6 App Shell + Sidebar"),
        ("2.9", "Task 8: Discovery Panel UI"),
        ("2.10", "Task 9: Settings Panel — All 5 Tabs"),
        ("2.11", "Task 10: Wire UI to Backend"),
        ("3", "Sub-project 2: LUT Upload Pipeline"),
        ("3.1", "bscpylgtv Upload API Reference"),
        ("3.2", "Task 1: LUT Data Models + .cube Parser"),
        ("3.3", "Task 2: LUT Uploader"),
        ("3.4", "Task 3: Dolby Vision Config Upload"),
        ("3.5", "Task 4: LUT Panel UI"),
        ("3.6", "Task 5: Hardware LUT Upload Tests"),
        ("3.7", "Task 6: Wire LUTPanel into MainWindow"),
        ("4", "Sub-project 3: Measurement Workflow"),
        ("4.1", "ArgyllCMS Primer"),
        ("4.2", "PGenerator 1.6 HTTP API"),
        ("4.3", "iTPG Internal Pattern Generator"),
        ("4.4", "Task 1: Measurement Data Models"),
        ("4.5", "Task 2: ArgyllCMS Subprocess Wrapper"),
        ("4.6", "Task 3: Pattern Generator Abstraction + iTPG"),
        ("4.7", "Task 4: PGenerator 1.6 HTTP Client"),
        ("4.8", "Task 5: Patch Sequences"),
        ("4.9", "Task 6: Measurement Session"),
        ("4.10", "Task 7: CGATS + JSON Storage"),
        ("4.11", "Task 8: LUT Generation from Measurements"),
        ("4.12", "Task 9: Measurement Panel UI"),
        ("4.13", "Task 10: Wire Measurement Workflow into MainWindow"),
        ("4.14", "Task 11: Hardware Integration Tests"),
    ]

    for number, title in toc_entries:
        add_toc_entry(doc, number, title)

    add_page_break(doc)


def parse_markdown_and_add(doc, md_text, section_prefix=""):
    """
    Parse a markdown string and add its content to the docx document.
    Handles: # headings, ## headings, ### headings, - lists, - [ ] checklists,
    ```code blocks```, **bold**, tables (|col|col|), blank lines.
    """
    lines = md_text.split('\n')
    in_code = False
    code_lang = ""
    code_lines = []
    in_table = False
    table_rows = []
    i = 0

    def flush_code():
        nonlocal code_lines, code_lang
        if not code_lines:
            return
        # Add a shaded code block
        for code_line in code_lines:
            p = doc.add_paragraph(style='Code')
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(code_line if code_line else " ")
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        code_lines = []
        code_lang = ""

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        # Filter out separator rows (---|--- pattern)
        data_rows = [r for r in table_rows if not re.match(r'^\s*[-|:\s]+$', r.replace('|', ''))]
        if not data_rows:
            table_rows = []
            return

        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            parsed.append(cells)

        if not parsed:
            table_rows = []
            return

        num_cols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=num_cols)
        tbl.style = 'Table Grid'

        for ri, row_data in enumerate(parsed):
            for ci in range(num_cols):
                cell_text = row_data[ci] if ci < len(row_data) else ""
                cell = tbl.rows[ri].cells[ci]
                if ri == 0:
                    set_cell_background(cell, '1565C0')
                    run = cell.paragraphs[0].add_run(cell_text)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                else:
                    bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
                    set_cell_background(cell, bg)
                    run = cell.paragraphs[0].add_run(cell_text)
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'

        doc.add_paragraph()
        table_rows = []

    while i < len(lines):
        line = lines[i]

        stripped_line = line.strip()

        # Code block handling:
        # - ` ```<lang>` while not in code → open code block
        # - ` ``` ` (bare, no lang) while in code → close code block
        # - ` ```<lang>` while already in code → treat as code content (nested block)
        if stripped_line.startswith('```'):
            lang_suffix = stripped_line[3:].strip()
            if not in_code:
                flush_table()
                in_code = True
                code_lang = lang_suffix
                code_lines = []
            elif lang_suffix == '':
                # Bare closing ```
                flush_code()
                in_code = False
            else:
                # ` ```python` / ` ```bash` etc. inside an already-open code block
                # — treat as literal code content, do NOT close the block
                code_lines.append(line)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table rows
        if line.strip().startswith('|') and '|' in line[1:]:
            flush_code()
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False

        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', stripped):
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        h_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            # Strip markdown bold/italic from heading text
            text = re.sub(r'\*+([^*]+)\*+', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            level = min(level, 3)
            h = doc.add_heading(text, level=level)
            if level == 1:
                h.runs[0].font.color.rgb = COLOR_HEADING1
                h.runs[0].font.size = Pt(16)
            elif level == 2:
                h.runs[0].font.color.rgb = COLOR_HEADING2
                h.runs[0].font.size = Pt(13)
            elif level == 3:
                h.runs[0].font.color.rgb = COLOR_HEADING3
                h.runs[0].font.size = Pt(11)
            i += 1
            continue

        # Checkbox list item (task step)
        cb_match = re.match(r'^-\s+\[( |x|X)\]\s+\*\*(.+?)\*\*\s*(.*)', stripped)
        if cb_match:
            checked = cb_match.group(1).lower() == 'x'
            label = cb_match.group(2)
            rest = cb_match.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            checkbox = '☑' if checked else '☐'
            run_cb = p.add_run(f"{checkbox}  ")
            run_cb.font.name = 'Calibri'
            run_cb.font.size = Pt(10)
            run_cb.font.color.rgb = COLOR_ACCENT if checked else RGBColor(0x66, 0x66, 0x66)
            run_label = p.add_run(label)
            run_label.font.name = 'Calibri'
            run_label.font.bold = True
            run_label.font.size = Pt(10)
            run_label.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            if rest:
                run_rest = p.add_run(f"  {rest}")
                run_rest.font.name = 'Calibri'
                run_rest.font.size = Pt(10)
            i += 1
            continue

        # Regular list item
        li_match = re.match(r'^[-*]\s+(.+)$', stripped)
        if li_match:
            content = li_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _add_inline_formatted_run(p, content)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        _add_inline_formatted_run(p, stripped)
        i += 1

    # Flush any remaining
    if in_code:
        flush_code()
    if in_table:
        flush_table()


def _add_inline_formatted_run(para, text):
    """Add a paragraph run handling **bold**, `code`, and plain text."""
    # Split on bold (**...**) and inline code (`...`)
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        else:
            if part:
                run = para.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)


def add_section_divider(doc, title, number):
    """Add a prominent section divider for each sub-project."""
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(f"Section {number}")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb = COLOR_HEADING1
    h.runs[0].font.size = Pt(20)
    h.paragraph_format.space_after = Pt(16)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_title_page(doc)
    add_toc(doc)

    section_numbers = [1, 2, 3, 4]

    for (section_title, md_file), sec_num in zip(PLAN_FILES, section_numbers):
        print(f"  Processing: {md_file.name} ...")
        md_text = md_file.read_text(encoding='utf-8')

        add_section_divider(doc, section_title, sec_num)
        doc.add_paragraph()

        # Remove the top-level H1 from the markdown (we use the section divider instead)
        lines = md_text.split('\n')
        if lines and lines[0].startswith('# '):
            md_text = '\n'.join(lines[1:])

        parse_markdown_and_add(doc, md_text)

    doc.save(OUTPUT_PATH)
    print(f"\nSaved: {OUTPUT_PATH}")
    print(f"Size: {OUTPUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    print("Generating LG OLED Calibration App Implementation Plan .docx ...")
    main()
